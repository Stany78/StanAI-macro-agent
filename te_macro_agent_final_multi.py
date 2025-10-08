#!/usr/bin/env python3
# te_macro_agent_auto_select_unified.py
# ES: finestra 60gg (invariato nello stile, con fix "%").ciao
# NOTIZIE: tassonomia unificata, dedup/merge PCE, fallback colore, cap/quote, cap rendimenti=1,
# fill-up a MIN_TARGET=12 (prima ≤N giorni, poi estendi a N+10 e fino a 30 giorni dal DB),
# ordinamento finale Colore ↓, Score ↓, Recency ↑.
# DB/Delta Mode (SQLite) con prune 60gg. Navigazione TE robusta (www + retry) e scroll via window.scrollBy.

import os, re, time, logging, unicodedata, sqlite3, hashlib
from difflib import SequenceMatcher
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional, Any

from dotenv import load_dotenv
from dateutil import parser as dtparser
from playwright.sync_api import sync_playwright
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ============= Setup & Config =============
script_dir = Path(__file__).parent.absolute()
env_path = script_dir / ".env"
load_dotenv(env_path)

def setup_logging(level=logging.INFO):
    logging.basicConfig(level=level,
                        format="%(asctime)s | %(levelname)s | %(message)s",
                        datefmt="%Y-%m-%d %H:%M:%S")

@dataclass
class Config:
    # Browser
    HEADLESS: bool = True
    NAV_TIMEOUT: int = 120_000
    SLOW_MO: int = 0

    # Output
    OUTPUT_DIR: str = str(script_dir)

    # Anthropic (ES + traduzioni/riassunti)
    ANTHROPIC_API_KEY: str = ""
    MODEL: str = "claude-3-haiku-20240307"
    MODEL_TEMP: float = 0.2
    MAX_TOKENS: int = 1500

    # Limiti testo
    SUMMARY_WORDS: int = 100
    ES_WORD_MIN: int = 300
    ES_WORD_PER_EXTRA_COUNTRY: int = 150

    # TradingEconomics
    BASE_URL: str = "https://www.tradingeconomics.com/stream?i=economy"

    # Menu paesi
    DEFAULT_COUNTRIES_MENU: List[str] = field(default_factory=lambda: [
        "United States", "Euro Area", "Germany", "United Kingdom",
        "Italy", "France", "China", "Japan", "Spain", "Netherlands", "European Union"
    ])

    # Finestra ES
    CONTEXT_DAYS: int = 60  # <– esteso a 60 giorni

    # ---- Delta Mode / DB ----
    DELTA_MODE: bool = True
    SCRAPE_HORIZON_DAYS: int = 7                # scraping ridotto per paesi già "caldi"
    WARMUP_NEW_COUNTRY_MIN: int = 40            # soglia elementi in DB per considerare "caldo"
    DB_PATH: str = str(script_dir / "news_cache.sqlite")
    PRUNE_DAYS: int = 60                        # <– prune DB a 60 giorni
    
    def __post_init__(self):
        # Ricarica la chiave dopo l'inizializzazione
        if not self.ANTHROPIC_API_KEY:
            self.ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
            if not self.ANTHROPIC_API_KEY:
                raise RuntimeError(f"ANTHROPIC_API_KEY non trovata in {env_path}")

# ============= Utilità tempo/recency =============
REL_RX = re.compile(r"\b(\d+)\s*(minute|hour|day|week|month)s?\s+ago\b", re.I)

def parse_age_days_from_text(time_text: str) -> Optional[float]:
    if not time_text: return None
    m = REL_RX.search(time_text or "")
    if m:
        q = int(m.group(1)); unit = m.group(2).lower()
        if unit.startswith("minute"): return q/1440.0
        if unit.startswith("hour"):   return q/24.0
        if unit.startswith("day"):    return float(q)
        if unit.startswith("week"):   return float(q)*7.0
        if unit.startswith("month"):  return 30.0
    try:
        dt = dtparser.parse(time_text, fuzzy=True)
        delta = datetime.now() - dt
        return max(0.0, delta.total_seconds()/86400.0)
    except Exception:
        return None

def recency_weight(days: Optional[float]) -> float:
    if days is None: return 0.6
    if days <= 1:    return 1.0
    if days <= 3:    return 0.85
    if days <= 7:    return 0.7
    if days <= 14:   return 0.5
    return 0.35

# ============= Country normalizer =============
COUNTRY_SYNONYMS = {
    "European Union": "Euro Area", "Eurozone": "Euro Area", "EU": "Euro Area",
    "U.S.": "United States", "US": "United States", "USA": "United States",
    "United States": "United States", "Euro Area": "Euro Area",
    "China": "China", "Germany": "Germany", "United Kingdom": "United Kingdom", "UK": "United Kingdom",
    "Italy": "Italy", "France": "France", "Japan": "Japan", "Spain": "Spain", "Netherlands": "Netherlands",
}
def normalize_country(name: str) -> str:
    n = (name or "").strip()
    return COUNTRY_SYNONYMS.get(n, n)

# ============= DB (SQLite) =============
def _norm_for_fp(s: str) -> str:
    if not s: return ""
    s = unicodedata.normalize("NFKD", s)
    s = re.sub(r"[–—\-:;,\.!?\(\)\[\]\{\}]+", " ", s.lower())
    return re.sub(r"\s+", " ", s).strip()

def _fp(it: dict) -> str:
    base = f"{it.get('country','')}|{_norm_for_fp(it.get('title',''))}|{_norm_for_fp((it.get('description') or '')[:200])}"
    return hashlib.sha1(base.encode("utf-8","ignore")).hexdigest()

def db_init(path: str):
    conn = sqlite3.connect(path)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS te_items (
            key TEXT PRIMARY KEY,
            country TEXT, title TEXT, description TEXT,
            time_text TEXT, importance INTEGER, category_raw TEXT,
            first_seen_ts REAL, last_seen_ts REAL
        )
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_country_seen ON te_items(country, last_seen_ts)")
    conn.commit()
    return conn

def db_upsert(conn, items: list):
    now = time.time()
    cur = conn.cursor()
    for it in items:
        k = _fp(it)
        cur.execute("""
            INSERT INTO te_items(key,country,title,description,time_text,importance,category_raw,first_seen_ts,last_seen_ts)
            VALUES (?,?,?,?,?,?,?, ?, ?)
            ON CONFLICT(key) DO UPDATE SET
              last_seen_ts=excluded.last_seen_ts,
              time_text=excluded.time_text,
              importance=excluded.importance
        """, (k, it.get("country",""), it.get("title",""), it.get("description",""),
              it.get("time",""), int(it.get("importance",0)), it.get("category_raw",""),
              now, now))
    conn.commit()

def db_count_by_country(conn, country: str) -> int:
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM te_items WHERE country=?", (country,))
    r = cur.fetchone()
    return int(r[0] or 0)

def db_load_recent(conn, countries: list, max_age_days: int = 60) -> list:
    """Carica fino a max_age_days dal DB; età = parse(time_text) se possibile, altrimenti delta da last_seen_ts"""
    if not countries: return []
    cutoff = time.time() - max_age_days*86400
    qs = ",".join("?"*len(countries))
    cur = conn.cursor()
    cur.execute(f"""
        SELECT country,title,description,time_text,importance,category_raw,last_seen_ts
        FROM te_items
        WHERE last_seen_ts >= ? AND country IN ({qs})
    """, [cutoff, *countries])
    out=[]
    now = time.time()
    for c,t,d,tt,imp,cat,seen in cur.fetchall():
        age_tt = parse_age_days_from_text(tt or "")
        age_db = max(0.0, (now - float(seen)) / 86400.0)
        age_days = age_tt if (age_tt is not None and age_tt <= 90.0) else age_db
        out.append({
            "country": c, "title": t or "", "description": d or "",
            "time": tt or "", "importance": int(imp or 0),
            "category_raw": cat or "", "age_days": age_days
        })
    return out

def db_prune(conn, max_age_days: int = 60):
    cutoff = time.time() - max_age_days*86400
    cur = conn.cursor()
    cur.execute("DELETE FROM te_items WHERE last_seen_ts < ?", (cutoff,))
    conn.commit()

# ============= Scraper TradingEconomics =============
class TEStreamScraper:
    def __init__(self, cfg: Config): self.cfg = cfg

    @staticmethod
    def _map_color_to_importance(class_text: str, style_text: str) -> int:
        t = (class_text or "").lower() + " " + (style_text or "").lower()
        if any(k in t for k in ["high-impact","impact-high","text-danger","badge-danger","#dc3545"]): return 3
        if any(k in t for k in ["text-primary","badge-primary","#0d6efd"]): return 2
        if any(k in t for k in ["text-info","badge-info","#0dcaf0"]): return 1
        return 0

    def scrape_30d(self, chosen_countries: List[str], max_days: int = 60) -> List[Dict[str, Any]]:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=self.cfg.HEADLESS, slow_mo=self.cfg.SLOW_MO,
                args=["--disable-blink-features=AutomationControlled","--disable-gpu"])
            context = browser.new_context(
                user_agent=("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                            "(KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36"),
                viewport={"width":1440,"height":900}, locale="en-US")
            def _block(route):
                try:
                    url = route.request.url
                    if any(x in url for x in [".png",".jpg",".jpeg",".gif",".webp",".svg",".woff",".woff2",".ttf",
                                              ".mp4",".avi",".webm",".css?","doubleclick","googletag","analytics"]):
                        return route.abort()
                except Exception: pass
                return route.continue_()
            try: context.route("**/*", _block)
            except Exception: pass

            page = context.new_page()

            # Navigazione robusta (www + retry)
            def safe_goto():
                candidates = []
                u = self.cfg.BASE_URL.strip()
                candidates.append(u)
                if "://www." not in u: candidates.append(u.replace("://", "://www.", 1))
                if "://www." in u: candidates.append(u.replace("://www.", "://", 1))
                last_err = None
                for url in dict.fromkeys(candidates):
                    try:
                        page.goto(url, wait_until="domcontentloaded", timeout=self.cfg.NAV_TIMEOUT)
                        return True
                    except Exception as e:
                        last_err = e
                        continue
                raise last_err if last_err else RuntimeError("Impossibile raggiungere TradingEconomics")

            try:
                safe_goto()
            except Exception as nav_err:
                logging.error("Navigazione fallita verso TradingEconomics: %s", nav_err)
                try: browser.close()
                except Exception: pass
                return []  # fallback al DB nel chiamante

            # Cookie
            for sel in ['#onetrust-accept-btn-handler', 'button:has-text("Accept")', '[class*="cookie"] button']:
                try:
                    b = page.locator(sel).first
                    if b and b.is_visible(): b.click(timeout=1000); page.wait_for_timeout(200); break
                except Exception: pass

            # Aspetta almeno una card
            try:
                page.wait_for_selector('li.te-stream-item, div.stream-item, article', timeout=10_000)
            except Exception:
                logging.warning("Nessuna card visibile entro 10s; continuo comunque.")

            # Scroll con early-stop (robusto, senza mouse.wheel)
            older_hits = 0
            for _ in range(100):
                try:
                    if page.is_closed():
                        break
                except Exception:
                    break

                try:
                    page.evaluate("window.scrollBy(0, 1600);")
                except Exception:
                    break

                page.wait_for_timeout(350)

                for sel in ['#stream-btn:has-text("More")','button:has-text("More")','a:has-text("More")']:
                    try:
                        btn = page.locator(sel).first
                        if btn and btn.is_visible():
                            btn.click()
                            page.wait_for_timeout(420)
                    except Exception:
                        pass

                try:
                    tails = page.evaluate("""() => {
                        const nodes = Array.from(document.querySelectorAll('li.te-stream-item, div.stream-item, article'));
                        return nodes.slice(-25).map(n => (n.querySelector('small')?.textContent || '').trim());
                    }""") or []
                    ages = []
                    for tx in tails:
                        a = parse_age_days_from_text(tx)
                        if a is not None:
                            ages.append(a)
                    if ages and min(ages) > max_days:
                        older_hits += 1
                    else:
                        older_hits = 0
                    if older_hits >= 2:
                        break
                except Exception:
                    pass

            # Estrazione (no filtro paese in JS)
            def _extract_all():
                return page.evaluate(
                    """() => {
                        const pick = (el, sels) => {
                          for (const s of sels) { const n = el.querySelector(s); if (n) { const t=(n.textContent||'').trim(); if (t) return t; } }
                          return '';
                        };
                        const selsCountry = [
                          'a.te-stream-country', '.te-stream-country',
                          'a[href*="/country/"]', 'a[href*="/countries/"]',
                          '.country a', '[data-entity="country"]', '[data-country]'
                        ];
                        const selsTitle = ['a.te-stream-title', 'h3', 'h2', 'a', '.te-title', 'strong'];
                        const nodes = Array.from(document.querySelectorAll('li.te-stream-item, div.stream-item, article'));
                        return nodes.map((el) => {
                          const country = pick(el, selsCountry);
                          const title   = pick(el, selsTitle);
                          const desc    = (el.querySelector('span.te-stream-item-description')?.textContent
                                           || el.querySelector('.desc')?.textContent
                                           || el.querySelector('p')?.textContent
                                           || el.textContent || '').trim();
                          const time    = (el.querySelector('small')?.textContent || '').trim();
                          const c_blob  = [el.getAttribute('class')||'',
                                           el.querySelector('.te-stream-impact')?.getAttribute('class')||'',
                                           el.querySelector('small')?.getAttribute('class')||'',
                                           el.querySelector('.te-stream-title')?.getAttribute('class')||''].join(' ');
                          const s_blob  = [el.getAttribute('style')||'',
                                           el.querySelector('.te-stream-impact')?.getAttribute('style')||'',
                                           el.querySelector('small')?.getAttribute('style')||''].join(' ');
                          const cat     = (el.querySelector('a.te-stream-category')?.textContent||'').trim();
                          return {country, title, description: desc.slice(0,2000), time_text: time, class_blob: c_blob, style_blob: s_blob, category_raw: cat};
                        });
                    }"""
                ) or []

            raw = _extract_all()
            if not raw:
                for _ in range(12):
                    try:
                        page.evaluate("window.scrollBy(0, 1800);")
                    except Exception:
                        break
                    page.wait_for_timeout(380)
                    for s in ['#stream-btn:has-text("More")','button:has-text("More")','a:has-text("More")']:
                        try:
                            btn = page.locator(s).first
                            if btn and btn.is_visible():
                                btn.click()
                                page.wait_for_timeout(420)
                        except Exception:
                            pass
                raw = _extract_all()

            browser.close()

        # Post-process & filtro Paesi in Python
        items: List[Dict[str, Any]] = []
        chosen_set = set(normalize_country(c) for c in chosen_countries)
        for r in raw:
            country_raw = (r.get("country") or "").strip()
            country = normalize_country(country_raw)
            if not country or country not in chosen_set:
                continue
            age_days = parse_age_days_from_text(r.get("time_text","")) or parse_age_days_from_text(r.get("description",""))
            if age_days is None or age_days > float(max_days):
                continue
            importance = self._map_color_to_importance(r.get("class_blob",""), r.get("style_blob",""))
            items.append({
                "country": country,
                "title": (r.get("title","") or "").strip(),
                "description": (r.get("description","") or "").strip(),
                "time": (r.get("time_text","") or "").strip(),
                "age_days": age_days,
                "importance": importance,
                "category_raw": (r.get("category_raw","") or "").strip(),
            })
        logging.info("Scraper: notizie raccolte (<=%sgg) paesi=%s -> %d",
                     max_days, ",".join(chosen_countries), len(items))
        return items

# ============= Classificazione & Score (NOTIZIE) =============
GDP_RX  = re.compile(r"\b(gdp|gross domestic product|gdp growth rate|growth)\b", re.I)
INFL_RX = re.compile(r"\b(cpi|pce|ppi|inflation|deflator|core|wage|wages|earnings)\b", re.I)
LAB_RX  = re.compile(r"\b(unemployment|jobless|payroll|nonfarm|claims)\b", re.I)
PMI_RX  = re.compile(r"\b(pmi|ism|ifo|nbs|caixin|sentiment|confidence|esi|zew|business\s+confidence|consumer\s+confidence|climate)\b", re.I)
HOU_RX  = re.compile(r"\b(housing|home\s+sales|mortgage|building\s+permits|housing\s+starts|existing\s+home|new\s+home|home\s+prices?|property\s+investment)\b", re.I)

# crescita include IP/retail/trade/ordini/profitti/Mx/credito/TSF/immatricolazioni/scorte
GROWTH_EXTRAS = re.compile(
    r"\b(industrial\s+production|produzione\s+industriale|retail\s+sales|vendite\s+al\s+dettaglio|"
    r"trade\s+balance|bilancia\s+commerciale|export|import|current\s+account|factory\s+orders|"
    r"industrial\s+profits?|profitti\s+industriali|loans?|credit(?:o|i)?|total\s+social\s+financing|tsf|"
    r"m[123]\b|money\s+supply|car\s+registrations?|immatricolazioni|inventor(?:y|ies)|scorte)\b",
    re.I
)

CATEGORY_WEIGHTS = {
    "crescita":0.30, "inflazione":0.30, "lavoro":0.25, "pmi":0.20, "housing":0.15, "rendimenti":0.12, "altro":0.10
}
RENDITI_KEYS = re.compile(r"\b(rendimento|rendimenti|treasury|decennale|yield|t-?note|t-?bond|curve)\b", re.I)
PMI_REGIONAL = re.compile(r"\b(Richmond|Kansas\s*City|Dallas|Philadelphia|Philly|Empire|New\s*York|NY\s*Fed|Chicago|Atlanta|Cleveland)\b", re.I)

# Preview/outlook
PREVIEW_PAT  = re.compile(
    r"\b("
    r"in\s+attesa|atteso|previsioni|previsto|"
    r"ahead of|before the|outlook|preview|in vista di|"
    r"settimana prossima|la\s+settimana\s+che\s+verrà|week\s+ahead|"
    r"questa\s+settimana|la\s+settimana\s+in\s+arrivo|"
    r"calendario\s+macro|agenda|cosa\s+aspettarsi"
    r")\b",
    re.IGNORECASE
)

def detect_category_from_text(text: str) -> str:
    t=(text or "").lower()
    if PMI_RX.search(t):       return "pmi"
    if GDP_RX.search(t) or GROWTH_EXTRAS.search(t): return "crescita"
    if INFL_RX.search(t):      return "inflazione"
    if LAB_RX.search(t):       return "lavoro"
    if HOU_RX.search(t):       return "housing"
    return "altro"

def has_numbers(text: str) -> bool:
    return bool(re.search(r"\d+(?:[.,]\d+)?\s*%?", text or ""))

def score_item(it: Dict[str, Any]) -> int:
    text = f"{it.get('title','')} {it.get('description','')}"
    cat = detect_category_from_text(text)
    base = CATEGORY_WEIGHTS.get(cat, 0.10)
    rb = 0.16 * recency_weight(it.get("age_days"))
    nb = 0.12 if has_numbers(text) else 0.0
    boost = 0.0
    if cat=="crescita" and (GDP_RX.search(text) or GROWTH_EXTRAS.search(text)): boost += 0.05
    if cat=="inflazione" and INFL_RX.search(text): boost += 0.04
    if cat=="lavoro" and LAB_RX.search(text): boost += 0.04
    if cat=="pmi": boost += 0.02
    score = int(max(0, min(100, round((base+rb+nb+boost)*100))))
    it["category_mapped"]=cat; it["score"]=score
    return score

# ============= ES (Executive Summary) – INVARIATO (solo fix "%") =============
def _normalize_spaces_in_perc(text: str) -> str:
    if not text: return text
    t = text
    t = re.sub(r"\s+%", "%", t)                                 # "2,4 %" => "2,4%"
    t = re.sub(r"%(?=[A-Za-zÀ-ÖØ-öø-ÿ])", "% ", t)              # "5,2%su" => "5,2% su"
    t = re.sub(r"(\d)[\s]+([.,])[\s]*(\d)", r"\1\2\3", t)       # "3, 4%" => "3,4%"
    t = re.sub(r"\s{2,}", " ", t)
    return t.strip()

def build_es_input_text(context_items: List[Dict[str, Any]]) -> str:
    def is_gdp(it: Dict[str, Any]) -> bool:
        t = (it.get("title","")+" "+it.get("description","")).lower()
        return bool(re.search(r"\b(gdp|gross domestic product|gdp growth rate)\b", t))
    gdp_items=[x for x in context_items if is_gdp(x)]
    other_items=[x for x in context_items if not is_gdp(x)]
    def sort_key(x): return (-int(x.get("importance",0)), x.get("age_days",999), (x.get("title","") or "")[:60])
    gdp_items.sort(key=sort_key); other_items.sort(key=sort_key)
    ordered = gdp_items + other_items
    blocks=[]
    for it in ordered:
        line=f"{it.get('country','')}: {it.get('title','')}. {it.get('description','')}".strip()
        blocks.append(_normalize_spaces_in_perc(line))
    return "\n".join(blocks)

PROMPT_ES = (
    "sei un analista macroeconomico e devi scrivere un report macroeconomico narrativo e coerente, "
    "mantenendo tutti i dati numerici forniti e senza introdurne di nuovi. "
    "Dai priorità alla lettura dei fondamentali: GDP/crescita, inflazione (CPI/PCE/PPI), "
    "mercato del lavoro (payroll/unemployment/claims) e PMI/ISM/IFO. "
    "Se sono presenti aggiornamenti recenti (ultimi 7 giorni) su questi indicatori, menzionali esplicitamente nella discussione. "
    "Non usare sezioni o elenchi puntati: costruisci 3–5 paragrafi fluidi con transizioni chiare. "
    "Integra i punti in un ragionamento unitario (domanda interna, momentum dell'inflazione, condizioni del lavoro, segnali anticipatori), "
    "collega i dati ai possibili passi di politica monetaria e allo scenario di crescita. "
    "Chiudi con un paragrafo conclusivo su prospettive e rischi (cosa può sorprendere al rialzo o al ribasso) coerenti con i dati citati."
)

def _strip_generic_intro(text: str) -> str:
    if not text:
        return text
    t = text.strip()
    # italiano: ecco un/il report|executive summary|resoconto...
    t = re.sub(
        r"^\s*(ecco\s+(un|il)\s+(report|executive\s*summary|resoconto(\s+macroeconomico)?))[^.\n:–—]{0,240}[:.\-–—]\s+",
        "", t, count=1, flags=re.IGNORECASE
    )
    # italiano: varianti comuni
    t = re.sub(r"^\s*(executive\s*summary\s*:?\s*)", "", t, flags=re.IGNORECASE)
    t = re.sub(r"^\s*(in\s+questo\s+report|di\s+seguito\s+un\s+report|presentiamo\s+un\s+report)\b[^.\n:–—]*[:.\-–—]\s+", "", t, count=1, flags=re.IGNORECASE)
    # inglese: here is/this report/we present...
    t = re.sub(
        r"^\s*(here\s+is\s+(a\s+)?(narrative\s+)?(macroeconomic\s+)?report|"
        r"here['']s\s+(a\s+)?report|in\s+this\s+report|below\s+is\s+(an\s+)?(executive\s+summary|report)|"
        r"we\s+present\s+(an\s+)?(executive\s+summary|report)|this\s+report\s+(provides|presents)|"
        r"executive\s+summary\s*:?)\s*[:.\-–—]\s+",
        "", t, count=1, flags=re.IGNORECASE
    )
    return t.strip()


class MacroSummarizer:
    def __init__(self, api_key: str, model: str, temp: float, max_tokens: int):
        import anthropic
        if not api_key: raise RuntimeError("ANTHROPIC_API_KEY non impostata nel .env")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model, self.temp, self.max_tokens = model, temp, max_tokens

    def _call_with_retry(self, messages, temperature, max_tokens, max_retries=5):
        """Chiama l'API con retry automatico in caso di rate limit"""
        for attempt in range(max_retries):
            try:
                return self.client.messages.create(
                    model=self.model,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    messages=messages
                )
            except Exception as e:
                error_str = str(e)
                # Verifica se è un rate limit error (429)
                if "rate_limit" in error_str.lower() or "429" in error_str:
                    if attempt < max_retries - 1:
                        # Backoff esponenziale: 2^attempt secondi
                        wait_time = 2 ** attempt
                        logging.warning(f"Rate limit raggiunto. Attendo {wait_time}s prima del retry {attempt+1}/{max_retries}...")
                        time.sleep(wait_time)
                        continue
                    else:
                        logging.error("Rate limit: tutti i tentativi falliti")
                        raise
                else:
                    # Se non è un rate
                    raise
            return None

    def executive_summary(self, context_items: List[Dict[str, Any]], cfg: Config, chosen_countries: List[str]) -> str:
        extra = max(0, len(chosen_countries)-1)
        target_words = cfg.ES_WORD_MIN + cfg.ES_WORD_PER_EXTRA_COUNTRY * extra
        content_text = build_es_input_text(context_items) if context_items else "Nessun contenuto."
        try:
            resp = self._call_with_retry(
                messages=[{"role":"user","content":f"{PROMPT_ES}\n\nLunghezza obiettivo: circa {target_words} parole.\n\nTESTO DA RIELABORARE:\n{content_text}"}],
                temperature=self.temp,
                max_tokens=min(cfg.MAX_TOKENS,1600)
            )
            text = (resp.content[0].text if resp and resp.content else "").strip()
            text = _normalize_spaces_in_perc(text)
            text = _strip_generic_intro(text)
            return text or "Executive Summary non disponibile."
        except Exception as e:
            logging.error("Errore ES: %s", e)
            return "Executive Summary non disponibile per errore di generazione."

    def summarize_item_it(self, item: Dict[str, Any], cfg: Config) -> str:
        title = (item.get("title","") or "").strip()
        desc  = (item.get("description","") or "").strip()
        country = (item.get("country","") or "").strip()
        text_in = f"TITOLO: {title}\nPAESE: {country}\nTESTO: {desc}"
        prompt = (
            "Scrivi un riassunto in ITALIANO della seguente notizia economica. "
            "Usa 100–120 parole, tono professionale e chiaro, senza elenchi puntati né sezioni. "
            "Mantieni tutti i dati numerici presenti nel testo (percentuali, livelli, variazioni) senza introdurne di nuovi. "
            "Evidenzia il messaggio macro principale e l'eventuale implicazione di policy. "
            "Inizia direttamente con il contenuto.\n\n"
            f"CONTENUTO:\n{text_in}"
        )
        try:
            resp = self._call_with_retry(
                messages=[{"role":"user","content":prompt}],
                temperature=min(self.temp,0.3),
                max_tokens=500
            )
            out = (resp.content[0].text if resp and resp.content else "").strip()
            out = _normalize_spaces_in_perc(out)
            return out
        except Exception as e:
            logging.error("Errore summarize_item_it: %s", e)
            return (title if title else "")[:180]

    def translate_it(self, text: str, cfg: Config) -> str:
        if not text: return ""
        prompt = "Traduci in ITALIANO il seguente titolo. Rispondi SOLO con il titolo tradotto, senza frasi introduttive.\n\n" + text
        try:
            resp = self._call_with_retry(
                messages=[{"role":"user","content":prompt}],
                temperature=min(self.temp,0.2),
                max_tokens=120
            )
            out = (resp.content[0].text if resp and resp.content else "").strip()
            out = _strip_translation_preambles(_normalize_spaces_in_perc(out))
            return out
        except Exception as e:
            logging.error("Errore translate_it: %s", e)
            return text

# ============= Pulizie testuali =============
def _strip_translation_preambles(text: str) -> str:
    if not text: return ""
    t = text.strip()
    for rx in [
        r"^\s*il\s+titolo\s+tradotto\s+in\s+italiano\s+è\s*:\s*",
        r"^\s*titolo\s+in\s+italiano\s*:\s*",
        r"^\s*traduzione\s+in\s+italiano\s*:\s*",
        r"^\s*ecco\s+la\s+traduzione\s+in\s+italiano.*?:\s*",
        r"^\s*ecco\s+la\s+traduzione\s*:\s*",
    ]: t = re.sub(rx, "", t, flags=re.IGNORECASE)
    return t.strip()

def _norm_text(s: str) -> str:
    if not s: return ""
    s = unicodedata.normalize("NFKD", s)
    s = re.sub(r"[–—\-:;,\.!?\(\)\[\]\{\}]+", " ", s.lower())
    return re.sub(r"\s+", " ", s).strip()

def _fix_glued_numbers(text: str) -> str:
    if not text: return ""
    t = _normalize_spaces_in_perc(text)
    t = re.sub(r"(\d)(?=[A-Za-zÀ-ÖØ-öø-ÿ])", r"\1 ", t)  # 3mln -> 3 mln
    return re.sub(r"\s{2,}", " ", t).strip()

# ============= Selezione NOTIZIE (regole + FILL-UP dal DB) =============
STOPWORDS_IT_EN = set("""
 a ad ai al alla alle agli all' allo anche ancora at attesa avanti be by con col coi come da dai dal dalla dalle dagli dallo
 dei del della delle dello degli di do due e ed entro era erano fin fino fra gli i il in into la le lo ma nel nella nelle nello negli non o od of on per piu più
 quasi se senza sia sono su sul sulla sulle sullo sugli the to tra tre un una uno vs vs. week settimana prossima next ahead before outlook preview in vista di
""".split())

def _token_key(s: str) -> List[str]:
    s = _norm_text(s)
    return [t for t in s.split() if t not in STOPWORDS_IT_EN and not re.fullmatch(r"\d+([.,]\d+)?%?", t)][:18]

def _similar(a: str, b: str) -> float:
    return SequenceMatcher(None, _norm_text(a), _norm_text(b)).ratio()

def _is_preview_or_calendar(it: Dict[str, Any]) -> bool:
    t = f"{it.get('title','')} {it.get('description','')}"
    return bool(PREVIEW_PAT.search(t))

def _is_result_data(it: Dict[str, Any]) -> bool:
    t = f"{it.get('title','')} {it.get('description','')}"
    tn = _norm_text(t)
    if PREVIEW_PAT.search(t): return False
    if re.search(r"\b(sara|saranno|verra|verranno|will\s+be|to\s+be\s+released|expected\s+to|is\s+expected|are\s+expected)\b", tn):
        return False
    if has_numbers(t) and re.search(r"\b(yoy|y/y|mom|m/m|qoq|q/q|annuo|mensile|trimestrale|bps|punti|%)\b", tn): return True
    if re.search(r"\b(e|e')\s+(salit[oaie]|sc[eè]s[oaie]|aumentat[oaie]|diminuit[oaie]|accelerat[oaie]|rallentat[oaie]|rivist[oaie]|stabilizzat[oaie]|pubblicat[oaie]|attestat[oaie])\b", tn): return True
    return False

def _theme_is_rendimenti(it: Dict[str, Any]) -> bool:
    t = f"{it.get('title','')} {it.get('description','')}"
    return bool(RENDITI_KEYS.search(t))

def _is_pce_headline(text: str) -> bool:
    tl = (text or "").lower()
    return ("pce" in tl) and ("core" not in tl) and (
        "price" in tl or "prices" in tl or "deflator" in tl or "index" in tl or
        "prezzo" in tl or "prezzi" in tl or "deflatore" in tl or "indice" in tl
    )

def _is_pce_core(text: str) -> bool:
    tl = (text or "").lower()
    is_core_word = ("core" in tl) or ("di base" in tl) or ("al netto" in tl)
    return ("pce" in tl) and is_core_word and (
        "price" in tl or "prices" in tl or "deflator" in tl or "index" in tl or
        "prezzo" in tl or "prezzi" in tl or "deflatore" in tl or "indice" in tl
    )

def _topic_signature(it: Dict[str, Any]) -> str:
    if _theme_is_rendimenti(it): return "yields"
    t = (it.get("title","")+" "+it.get("description","") or "").lower()
    if _is_pce_headline(t) or _is_pce_core(t): return "pce"
    if "gdp" in t or "gross domestic product" in t or GROWTH_EXTRAS.search(t): return "gdp"
    if any(k in t for k in ["nonfarm","jobless","unemployment","payroll","claims"]): return "labour"
    if any(k in t for k in ["pmi","ism","ifo","nbs","caixin","sentiment","confidence","esi","zew","s&p global"]): return "pmi"
    if any(k in t for k in ["housing","home sales","mortgage","building permits","housing starts","existing home","new home","home prices","property investment"]): return "housing"
    key = " ".join(_token_key(it.get("title",""))[:6])
    return key or it.get("title","")[:30]

def _day_bucket(it: Dict[str, Any]) -> int:
    d = it.get("age_days")
    try: return int(d) if d is not None else 999
    except Exception: return 999

def _enrich_items(cands: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out=[]
    for it in cands:
        it = dict(it)  # copia
        score_item(it)
        t = f"{it.get('title','')} {it.get('description','')}"
        it["is_preview"]   = _is_preview_or_calendar(it)
        it["is_result"]    = _is_result_data(it)
        it["topic_sig"]    = _topic_signature(it)
        it["pce_headline"] = _is_pce_headline(t)
        it["pce_core"]     = _is_pce_core(t)
        it["day_bucket"]   = _day_bucket(it)
        if _theme_is_rendimenti(it):
            it["category_mapped"] = "rendimenti"
        # Fallback impatto/colore quando manca il flag TE
        if int(it.get("importance", 0)) == 0:
            cat = it.get("category_mapped","altro")
            if it["is_result"] and cat in {"crescita","inflazione","lavoro"}:
                it["importance"] = 3
            elif it["is_result"] and cat in {"pmi","housing"}:
                it["importance"] = 2
            elif not it["is_result"] and cat in {"crescita","inflazione","lavoro","pmi","housing"}:
                it["importance"] = 2
            else:
                it["importance"] = 1 if not it["is_preview"] else 0
        out.append(it)
    return out

def _group_and_clean(arr: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for it in arr: grouped.setdefault(it["topic_sig"], []).append(it)
    cleaned: List[Dict[str, Any]] = []
    for sig, items in grouped.items():
        results = [x for x in items if x["is_result"]]
        base = results if results else items
        if not base: continue
        # Merge PCE headline+core
        if sig == "pce":
            base_sorted = sorted(base, key=lambda i: (i.get("age_days",999), -int(i.get("score",0))))
            head = next((x for x in base_sorted if x.get("pce_headline")), None)
            core = next((x for x in base_sorted if x.get("pce_core")), None)
            if head and core:
                merged = dict(head)
                desc = (head.get("description","") or "")
                more = core.get("description","") or ""
                if more and SequenceMatcher(None, _norm_text(more), _norm_text(desc)).ratio() < 0.92:
                    desc = (desc + "\n" + more).strip()
                merged["description"] = desc
                merged["merged_from"] = ["pce_headline","pce_core"]
                cleaned.append(merged)
                continue
        # Dedup per titolo
        kept: List[Dict[str, Any]] = []
        for cand in sorted(base, key=lambda i: (-int(i.get("score",0)), i.get("age_days",999))):
            if not kept: kept.append(cand); continue
            if any(SequenceMatcher(None, _norm_text(cand.get("title","")), _norm_text(k.get("title",""))).ratio() >= 0.92 for k in kept):
                continue
            kept.append(cand)
        cleaned.extend(kept)
    return cleaned

def _filter_nonreds_base(nonreds: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    filtered=[]
    for it in nonreds:
        if it.get("category_mapped") == "pmi" and PMI_REGIONAL.search(f"{it.get('title','')} {it.get('description','')}"):
            continue
        if it.get("category_mapped") == "altro":
            if not has_numbers(f"{it.get('title','')} {it.get('description','')}"):
                continue
        filtered.append(it)
    return filtered

def _take_nonreds_with_caps(nonreds_sorted, cap_total=16, cap_per_day=6, quota=None, yields_cap=1, already=None):
    if quota is None:
        quota = {"crescita":3, "inflazione":3, "lavoro":3, "pmi":2, "housing":1}
    already = already or []

    # stato iniziale derivato da 'already'
    from collections import defaultdict
    per_day = defaultdict(int)
    yields_used = 0
    for s in already:
        d = s.get("day_bucket", 999)
        per_day[d] += 1
        if (s.get("topic_sig") == "yields") or (s.get("category_mapped") == "rendimenti"):
            yields_used += 1

    counts_cat = {k: 0 for k in quota}
    for s in already:
        cat = s.get("category_mapped")
        if cat in counts_cat:
            counts_cat[cat] += 1

    chosen = []
    def _dup(cand):
        return any(SequenceMatcher(None, _norm_text(cand.get("title","")), _norm_text(s.get("title",""))).ratio() >= 0.92 for s in already+chosen)

    count_total = 0
    for it in nonreds_sorted:
        if count_total >= cap_total:
            break
        if _dup(it):
            continue
        day = it.get("day_bucket", 999)
        if per_day[day] >= cap_per_day:
            continue
        if (it.get("topic_sig")=="yields" or it.get("category_mapped")=="rendimenti"):
            if yields_used >= yields_cap:
                continue
        cat = it.get("category_mapped","altro")
        if cat in counts_cat and counts_cat[cat] >= quota[cat]:
            continue
        if cat not in counts_cat and cat not in {"rendimenti","altro"}:
            continue

        chosen.append(it)
        count_total += 1
        per_day[day] += 1
        if cat in counts_cat:
            counts_cat[cat] += 1
        if (it.get("topic_sig")=="yields" or it.get("category_mapped")=="rendimenti"):
            yields_used += 1

    return chosen


def _sort_final_key(it: Dict[str, Any]):
    return (-int(it.get("importance",0)), -int(it.get("score",0)), it.get("age_days",999))

def build_selection(items_ctx: List[Dict[str, Any]], days: int, cfg: Config,
                    expand1_days: int = 10, expand2_days: int = 30) -> List[Dict[str, Any]]:
    """
    Selezione: pool primario ≤N giorni; se <12 o mancano categorie core,
    estendi fino a N+10 e poi fino a 30 giorni (dal DB) rispettando cap/quote.
    """
    MIN_TARGET = 12
    # --- POOL 0: ≤ N giorni
    pool0 = [x for x in items_ctx if x.get("age_days") is not None and x["age_days"] <= float(days)]
    if not pool0: pool0=[]
    pool0 = _enrich_items(pool0)

    reds0    = [x for x in pool0 if int(x.get("importance",0)) == 3]
    nonreds0 = [x for x in pool0 if int(x.get("importance",0)) != 3]
    reds0    = _group_and_clean(reds0)
    nonreds0 = _group_and_clean(nonreds0)
    nonreds0 = _filter_nonreds_base(nonreds0)
    nonreds0.sort(key=lambda i: (-int(i.get("score",0)), i.get("age_days",999)))

    final_list: List[Dict[str, Any]] = []
    final_list.extend(reds0)
    chosen_nonreds0 = _take_nonreds_with_caps(nonreds0, already=final_list)
    final_list.extend(chosen_nonreds0)

    # copertura categorie core (crescita, inflazione, lavoro, pmi)
    macro_core = ["crescita","inflazione","lavoro","pmi"]
    present_core = set([it.get("category_mapped") for it in final_list if it.get("category_mapped") in macro_core and it.get("is_result")])

    need_fill = (len(final_list) < MIN_TARGET) or any(c not in present_core for c in macro_core)

    # --- helper per fill-up con priorità risultato > preview > altro con numeri
    def _fill_from_pool(pool_all: List[Dict[str, Any]], min_age_exclusive: float, max_age_inclusive: float, target: int, ensure_core: bool):
        nonlocal final_list
        pool = [x for x in pool_all if x.get("age_days") is not None and min_age_exclusive < x["age_days"] <= max_age_inclusive]
        if not pool: return
        pool = _enrich_items(pool)
        reds    = [x for x in pool if int(x.get("importance",0)) == 3]
        nonreds = [x for x in pool if int(x.get("importance",0)) != 3]
        reds    = _group_and_clean(reds)
        nonreds = _group_and_clean(nonreds)
        nonreds = _filter_nonreds_base(nonreds)

        # 1) Se serve copertura core: prendo prima risultati per core mancanti
        if ensure_core:
            current_core = set([it.get("category_mapped") for it in final_list if it.get("category_mapped") in macro_core and it.get("is_result")])
            missing = [c for c in macro_core if c not in current_core]
            if missing:
                cands_res = [x for x in (reds+nonreds) if x.get("is_result") and x.get("category_mapped") in missing]
                cands_res.sort(key=lambda i: (-int(i.get("score",0)), i.get("age_days",999)))
                for c in cands_res:
                    if len(final_list) >= target: break
                    # applica caps come per non-rossi
                    chosen = _take_nonreds_with_caps([c], already=final_list)
                    if chosen:
                        final_list.extend(chosen)
                        current_core.add(chosen[0].get("category_mapped"))

            # se ancora mancano core, prendo preview
            missing = [c for c in macro_core if c not in current_core]
            if missing and len(final_list) < target:
                cands_prev = [x for x in (reds+nonreds) if not x.get("is_result") and x.get("category_mapped") in missing]
                cands_prev.sort(key=lambda i: (-int(i.get("score",0)), i.get("age_days",999)))
                for c in cands_prev:
                    if len(final_list) >= target: break
                    chosen = _take_nonreds_with_caps([c], already=final_list)
                    if chosen:
                        final_list.extend(chosen)
                        current_core.add(chosen[0].get("category_mapped"))

        # 2) Se sotto target, prendo non-rossi migliori rispettando caps
        if len(final_list) < target:
            nonreds.sort(key=lambda i: (-int(i.get("score",0)), i.get("age_days",999)))
            chosen_more = _take_nonreds_with_caps(nonreds, already=final_list)
            for c in chosen_more:
                if len(final_list) >= target: break
                final_list.append(c)

        # 3) Se ancora sotto, prendo "altro con numeri"
        if len(final_list) < target:
            altri = [x for x in nonreds if x.get("category_mapped")=="altro" and has_numbers(f"{x.get('title','')} {x.get('description','')}")]
            altri.sort(key=lambda i: (-int(i.get("score",0)), i.get("age_days",999)))
            for c in altri:
                if len(final_list) >= target: break
                chosen = _take_nonreds_with_caps([c], already=final_list)
                if chosen: final_list.extend(chosen)

    if need_fill:
        # Fase 1: estendi a N + 10 (entro 30gg max)
        up1 = min(days + expand1_days, 30)
        _fill_from_pool(items_ctx, min_age_exclusive=days, max_age_inclusive=up1, target=MIN_TARGET, ensure_core=True)

    if len(final_list) < MIN_TARGET:
        # Fase 2: estendi fino a 30 gg
        if days < 30:
            _fill_from_pool(items_ctx, min_age_exclusive=min(days+expand1_days,30), max_age_inclusive=30, target=MIN_TARGET, ensure_core=True)

    # Ordinamento finale
    final_list.sort(key=_sort_final_key)
    return final_list

# ============= Report DOCX =============
def trim_words(text: str, max_words: int) -> str:
    if not text: return ""
    ws = text.strip().split()
    if len(ws) <= max_words: return text.strip()
    cut = " ".join(ws[:max_words]).rstrip()
    if "." in cut:
        tmp = ".".join(cut.split(".")[:-1]).strip()
        if tmp: return tmp + "."
    return cut

def save_report(filename: str, es_text: str, selection: List[Dict[str, Any]],
                countries: List[str], days: int, context_count: int, output_dir: str) -> str:
    doc = Document()
    title = doc.add_heading('MACRO MARKETS ANALYSIS – Selezione Automatica', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Executive Summary basato su {context_count} notizie (DB+stream) degli ultimi 60 giorni.")
    doc.add_paragraph(f"Paesi considerati (60gg): {', '.join(countries)}")
    doc.add_paragraph(f"Selezione mostrata: ultimi {days} giorni (ordinata per impatto, score, recency)")
    doc.add_paragraph(f"Data Report: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("Fonte: TradingEconomics (Stream) + cache locale")
    doc.add_paragraph("_" * 60)

    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    doc.add_paragraph(es_text or "(non disponibile)")
    doc.add_paragraph("_" * 60)

    doc.add_heading('NOTIZIE SELEZIONATE', level=1)
    by_country: Dict[str, List[Dict[str, Any]]] = {}
    for it in selection:
        by_country.setdefault(it.get("country","Unknown"), []).append(it)
    for country, arr in by_country.items():
        doc.add_heading(country.upper(), level=2)
        for i, it in enumerate(arr, 1):
            head_raw = it.get('title_it') or it.get('title','')
            head = f"{i}. {_fix_glued_numbers(_strip_translation_preambles(head_raw))}"
            p = doc.add_paragraph(); p.add_run(head).bold = True

            meta = []
            if it.get('time'): meta.append(f"⏰ {it['time']}")
            if it.get('category_mapped'): meta.append(f"🏷 {it['category_mapped']}")
            if it.get('score') is not None: meta.append(f"⭐ {it['score']}")
            if it.get('importance') is not None: meta.append(f"🎯 {it['importance']}")
            if it.get('merged_from') and it.get('topic_sig','') == 'pce':
                meta.append("🧩 pce: headline+core")
            if meta: doc.add_paragraph("   " + "  ·  ".join(meta))

            text_it = it.get('summary_it') or it.get('description','')
            text_it = _fix_glued_numbers(_strip_translation_preambles(text_it))
            doc.add_paragraph(trim_words(text_it, 120))
            doc.add_paragraph("")

    out_dir = Path(output_dir); out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename
    doc.save(out_path)
    logging.info("Report salvato: %s", out_path)
    return str(out_path)

# ============= Input & Main =============
def prompt_days() -> int:
    while True:
        s = input("Giorni da mostrare nella SELEZIONE (1-30): ").strip()
        if s.isdigit():
            d = int(s)
            if 1 <= d <= 30:
                return d
        print("Valore non valido. Inserisci un numero tra 1 e 30.")

def prompt_countries(default_menu: List[str]) -> List[str]:
    mapping = {"1":"United States","2":"Euro Area","3":"Germany","4":"United Kingdom",
               "5":"Italy","6":"France","7":"China","8":"Japan","9":"Spain","10":"Netherlands","11":"European Union"}
    print("\nSeleziona i Paesi (invio per default: United States, Euro Area):")
    print("  1) United States   2) Euro Area   3) Germany   4) United Kingdom   5) Italy")
    print("  6) France          7) China       8) Japan     9) Spain            10) Netherlands")
    print("  11) European Union")
    s = input("Inserisci codici separati da virgola (es. 1,2,4) oppure premi invio per default: ").strip()
    if not s: return ["United States","Euro Area"]
    chosen=[]
    for tok in s.split(","):
        k = tok.strip()
        if k in mapping and mapping[k] not in chosen:
            chosen.append(mapping[k])
    if not chosen: return ["United States","Euro Area"]
    return ["Euro Area" if x=="European Union" else x for x in chosen]

def get_or_prompt_anthropic_key(cfg: Config) -> str:
    if cfg.ANTHROPIC_API_KEY: return cfg.ANTHROPIC_API_KEY
    print("\n🔑 Inserisci la tua API Key Anthropic (verrà salvata in .env):")
    key = input("API Key: ").strip()
    if not key: raise RuntimeError("API Key obbligatoria per generare ES.")
    try:
        with open(env_path,"a",encoding="utf-8") as f: f.write(f"\nANTHROPIC_API_KEY={key}\n")
        os.environ["ANTHROPIC_API_KEY"] = key
    except Exception as e:
        logging.warning("Impossibile salvare .env: %s", e)
    return key

def main():
    setup_logging(logging.INFO)
    cfg = Config()
    print("="*100)
    print("🏦 TE Macro Agent – ES 60gg (invariato) | Delta Mode (DB) + Selezione con Fill-Up")
    print("="*100)

    # API Key
    try:
        get_or_prompt_anthropic_key(cfg)
    except Exception as e:
        print(f"\n❌ {e}")
        return

    # Input
    selection_days = prompt_days()
    chosen_countries = prompt_countries(cfg.DEFAULT_COUNTRIES_MENU)
    chosen_countries = ["Euro Area" if x=="European Union" else x for x in chosen_countries]
    print(f"\n▶ Contesto ES: {cfg.CONTEXT_DAYS} giorni | Selezione: {selection_days} giorni | Paesi: {', '.join(chosen_countries)}")

    scraper = TEStreamScraper(cfg)

    # ==== Pipeline dati con DB (Delta Mode) ====
    items_ctx: List[Dict[str, Any]] = []
    if cfg.DELTA_MODE:
        conn = db_init(cfg.DB_PATH)

        warm, fresh = [], []
        for c in chosen_countries:
            cnt = db_count_by_country(conn, c)
            (warm if cnt >= cfg.WARMUP_NEW_COUNTRY_MIN else fresh).append(c)

        items_new = []
        if fresh:
            logging.info("Warm-up paesi nuovi (<=%sgg): %                    s", cfg.CONTEXT_DAYS, ", ".join(fresh))
            items_new += scraper.scrape_30d(fresh, max_days=cfg.CONTEXT_DAYS)
        if warm:
            logging.info("Delta scrape paesi noti (<=%sgg): %s", cfg.SCRAPE_HORIZON_DAYS, ", ".join(warm))
            items_new += scraper.scrape_30d(warm, max_days=min(cfg.SCRAPE_HORIZON_DAYS, cfg.CONTEXT_DAYS))

        if items_new:
            db_upsert(conn, items_new)
            db_prune(conn, max_age_days=cfg.PRUNE_DAYS)

        items_ctx = db_load_recent(conn, chosen_countries, max_age_days=cfg.CONTEXT_DAYS)

        # fallback: base scarsa → scrape completo finestra ES
        if len(items_ctx) < 20:
            logging.info("Base DB scarsa (%d). Fallback scrape <=%sgg per tutti i paesi scelti.", len(items_ctx), cfg.CONTEXT_DAYS)
            items_all = scraper.scrape_30d(chosen_countries, max_days=cfg.CONTEXT_DAYS)
            if items_all:
                db_upsert(conn, items_all)
                items_ctx = db_load_recent(conn, chosen_countries, max_age_days=cfg.CONTEXT_DAYS)
    else:
        items_ctx = scraper.scrape_30d(chosen_countries, max_days=cfg.CONTEXT_DAYS)

    if not items_ctx:
        print("\n❌ Nessuna notizia disponibile (entro la finestra).")
        return

    # ==== ES (60gg) – invariato (solo fix "%") ====
    summarizer = MacroSummarizer(cfg.ANTHROPIC_API_KEY, cfg.MODEL, cfg.MODEL_TEMP, cfg.MAX_TOKENS)
    es_text = summarizer.executive_summary(items_ctx, cfg, chosen_countries)

    # ==== Selezione (ultimi N giorni) + Fill-Up dal DB ====
    selection_items = build_selection(items_ctx, selection_days, cfg, expand1_days=10, expand2_days=30)

    # ==== Traduzione titoli + riassunti IT con delay tra le chiamate ====
    for i, it in enumerate(selection_items):
        # Aggiungi un piccolo delay tra le richieste (tranne la prima)
        if i > 0:
            time.sleep(0.5)  # Pausa di 500ms tra una richiesta e l'altra
        
        try:
            it["title_it"] = summarizer.translate_it(it.get("title", ""), cfg)
        except Exception as e:
            logging.warning("Titolo non tradotto: %s", e)
            it["title_it"] = it.get("title", "")
        
        try:
            it["summary_it"] = summarizer.summarize_item_it(it, cfg)
        except Exception as e:
            logging.warning("Riassunto IT non disponibile: %s", e)
            it["summary_it"] = (it.get("description","") or "")

    # ==== Report ====
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"MacroAnalysis_AutoSelect_{selection_days}days_{ts}.docx"
    out_path = save_report(filename, es_text, selection_items, chosen_countries, selection_days, len(items_ctx), cfg.OUTPUT_DIR)

    print("\n" + "="*80)
    print("✅ COMPLETATO")
    print("="*80)
    print(f"• Report salvato in: {out_path}")
    print(f"• Notizie totali (DB, ultimi {cfg.CONTEXT_DAYS} gg): {len(items_ctx)}")
    print(f"• Notizie selezionate (ultimi {selection_days} gg + fill-up): {len(selection_items)}")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n⚠️ Interrotto dall'utente")
    except Exception as e:
        logging.exception("ERRORE FATALE: %s", e)
        print(f"\n❌ ERRORE: {e}")                         